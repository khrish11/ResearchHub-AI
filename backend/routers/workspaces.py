from fastapi import APIRouter, Depends, HTTPException, Response
from typing import Any, Dict, List, Optional, Tuple
from pydantic import BaseModel, ConfigDict, Field
import json
import io
import re
import textwrap
import os
from collections import Counter
from datetime import datetime, timezone

from utils.groq_client import client as groq_client, model_config
from utils.firebase_storage import download_bytes, storage_is_configured, upload_bytes

from repositories.research import User, Workspace, Paper, UserSessionState, WorkspaceDocument
from repositories import ResearchRepository, get_research_repository
from routers.auth import get_current_user

router = APIRouter(prefix="/workspaces", tags=["workspaces"])


class WorkspaceCreate(BaseModel):
    name: str
    description: Optional[str] = None


class WorkspaceOut(BaseModel):
    id: int
    name: str
    description: Optional[str]

    model_config = ConfigDict(from_attributes=True)


class PaperOut(BaseModel):
    id: int
    title: str
    authors: str
    abstract: str
    url: Optional[str] = None
    doi: Optional[str] = None
    bibcode: Optional[str] = None
    source: Optional[str] = None
    pdf_url: Optional[str] = None
    institutional_url: Optional[str] = None
    access_type: Optional[str] = None
    full_text_available: Optional[bool] = None

    model_config = ConfigDict(from_attributes=True)


class ChatOut(BaseModel):
    id: int
    message: str
    response: str

    model_config = ConfigDict(from_attributes=True)


class WorkspaceDetail(BaseModel):
    id: int
    name: str
    description: Optional[str]
    papers: List[PaperOut]
    chats: List[ChatOut]

    model_config = ConfigDict(from_attributes=True)


class ResearchReportRequest(BaseModel):
    topic: Optional[str] = None
    paper_ids: Optional[List[int]] = None
    depth: Optional[str] = "balanced"  # quick | balanced | deep
    focus_mode: Optional[str] = "broad"  # broad | methods | applications | risks


class SessionStateUpdate(BaseModel):
    page_path: Optional[str] = None
    workspace_id: Optional[int] = None
    last_query: Optional[str] = None
    draft_text: Optional[str] = None
    extra: Optional[Dict[str, Any]] = None


class SessionStateOut(BaseModel):
    page_path: str
    workspace_id: Optional[int] = None
    last_query: Optional[str] = None
    draft_text: Optional[str] = None
    extra: Dict[str, Any] = Field(default_factory=dict)
    updated_at: Optional[str] = None


class DocspaceDocumentUpdate(BaseModel):
    title: Optional[str] = None
    content: str = Field(default="", max_length=200000)


class DocspaceDocumentOut(BaseModel):
    workspace_id: int
    title: str
    content: str
    version: int
    updated_at: Optional[str] = None


class WorkspaceFileOut(BaseModel):
    id: int
    workspace_id: int
    user_id: int
    paper_id: Optional[int] = None
    kind: str
    filename: str
    storage_bucket: str
    storage_path: str
    download_url: Optional[str] = None
    content_type: Optional[str] = None
    size_bytes: int
    created_at: Optional[str] = None


def _normalize_report_depth(value: Optional[str]) -> str:
    depth = (value or "balanced").strip().lower()
    if depth not in {"quick", "balanced", "deep"}:
        return "balanced"
    return depth


def _normalize_focus_mode(value: Optional[str]) -> str:
    mode = (value or "broad").strip().lower()
    if mode not in {"broad", "methods", "applications", "risks"}:
        return "broad"
    return mode


def _safe_filename(text: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9_-]+", "-", (text or "").strip().lower()).strip("-")
    return slug or "research-report"


def _safe_storage_filename(filename: str, fallback: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "-", (filename or "").strip()).strip("-")
    return cleaned or fallback


def _backend_base_url() -> str:
    return (os.getenv("BACKEND_URL") or "http://127.0.0.1:8010").rstrip("/")


def _normalize_page_path(value: Optional[str]) -> str:
    path = (value or "/home").strip()
    if not path.startswith("/"):
        path = "/" + path
    return path[:240]


def _decode_extra_json(raw: Optional[str]) -> Dict[str, Any]:
    if not raw:
        return {}
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        return {}
    return {}


def _state_to_out(state: Optional[UserSessionState]) -> SessionStateOut:
    if not state:
        return SessionStateOut(page_path="/home", extra={})
    return SessionStateOut(
        page_path=_normalize_page_path(state.page_path),
        workspace_id=state.workspace_id,
        last_query=state.last_query,
        draft_text=state.draft_text,
        extra=_decode_extra_json(state.extra_json),
        updated_at=(
            state.updated_at.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")
            if state.updated_at
            else None
        ),
    )


def _iso_utc(value: Optional[datetime]) -> Optional[str]:
    if not value:
        return None
    return value.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")


def _persist_workspace_file_if_configured(
    *,
    repo: ResearchRepository,
    workspace_id: int,
    user_id: int,
    kind: str,
    filename: str,
    content: bytes,
    content_type: str,
) -> Optional[object]:
    if not storage_is_configured():
        return None
    safe_name = _safe_storage_filename(filename, fallback="artifact.bin")
    storage_path = (
        f"workspace-files/{user_id}/{workspace_id}/exports/{kind}/{safe_name}"
    )
    uploaded = upload_bytes(
        storage_path=storage_path,
        data=content,
        content_type=content_type,
        metadata={
            "workspace_id": str(workspace_id),
            "user_id": str(user_id),
            "kind": kind,
            "filename": safe_name,
        },
    )
    record = repo.create_workspace_file(
        workspace_id=workspace_id,
        user_id=user_id,
        kind=kind,
        filename=safe_name,
        storage_bucket=uploaded.bucket,
        storage_path=uploaded.path,
        content_type=uploaded.content_type,
        size_bytes=uploaded.size_bytes,
    )
    record.download_url = (
        f"{_backend_base_url()}/workspaces/{workspace_id}/files/{record.id}/download"
    )
    repo.save(record)
    return record


from utils.text_utils import (
    STOP_WORDS as _STOP_WORDS,
    tokenize as _tokenize_for_rank,
    extract_keywords as _extract_keywords_shared,
)


def _extract_keywords(text: str, top_n: int = 12) -> List[str]:
    return _extract_keywords_shared(text, max_keywords=top_n)


def _paper_relevance_score(topic: str, paper: Paper) -> int:
    query_set = set(_tokenize_for_rank(topic))
    if not query_set:
        return 0
    title_set = set(_tokenize_for_rank(paper.title or ""))
    abstract_set = set(_tokenize_for_rank(paper.abstract or ""))
    title_hits = len(query_set.intersection(title_set))
    abstract_hits = len(query_set.intersection(abstract_set))
    return (title_hits * 4) + (abstract_hits * 2)


def _select_ranked_papers(
    topic: str, papers: List[Paper], limit: int = 40
) -> List[Paper]:
    ranked = sorted(
        papers,
        key=lambda paper: (
            _paper_relevance_score(topic, paper),
            len((paper.abstract or "")),
            len((paper.title or "")),
        ),
        reverse=True,
    )
    return ranked[:limit]


def _paper_primary_link(paper: Paper) -> str:
    url = (paper.url or "").strip()
    doi = (getattr(paper, "doi", "") or "").strip()
    if url:
        return url
    if doi:
        clean_doi = (
            doi.replace("https://doi.org/", "").replace("http://doi.org/", "").strip()
        )
        return f"https://doi.org/{clean_doi}" if clean_doi else ""
    return ""


def _build_context(topic: str, selected_papers: List[Paper]) -> str:
    lines = [
        f"Topic: {topic}",
        "",
        "Workspace papers (use references like 'Paper 1', 'Paper 2'):",
    ]
    for idx, paper in enumerate(selected_papers, start=1):
        abstract = (paper.abstract or "").replace("\n", " ").strip()
        if len(abstract) > 1200:
            abstract = abstract[:1200] + "..."
        doi = getattr(paper, "doi", "") or "N/A"
        url = _paper_primary_link(paper) or "N/A"
        lines.append(
            (
                f"Paper {idx} Title: {paper.title}\n"
                f"Authors: {paper.authors}\n"
                f"DOI: {doi}\n"
                f"URL: {url}\n"
                f"Abstract: {abstract or 'No abstract available.'}\n"
            )
        )
    return "\n".join(lines)[:36000]


def _paper_links_markdown(selected_papers: List[Paper], max_items: int = 16) -> str:
    rows: List[str] = []
    for idx, paper in enumerate(selected_papers[:max_items], start=1):
        title = (paper.title or f"Paper {idx}").strip()
        link = _paper_primary_link(paper)
        if link:
            rows.append(f"- Paper {idx}: [{title}]({link})")
        else:
            rows.append(f"- Paper {idx}: {title} (Link unavailable)")
    return "\n".join(rows) or "- Paper links unavailable."


def _paper_links_data(
    selected_papers: List[Paper], max_items: int = 24
) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    for idx, paper in enumerate(selected_papers[:max_items], start=1):
        doi = (getattr(paper, "doi", "") or "").strip()
        link = _paper_primary_link(paper)
        items.append(
            {
                "paper": f"Paper {idx}",
                "index": idx,
                "title": (paper.title or f"Paper {idx}").strip(),
                "url": link or None,
                "doi": doi or None,
                "link_available": bool(link),
            }
        )
    return items


def _extract_mindmap_section(markdown_text: str) -> List[str]:
    lines = (markdown_text or "").splitlines()
    in_section = False
    section: List[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped.lower().startswith("## mindmap"):
            in_section = True
            continue
        if in_section and re.match(r"^\s*##\s+", line):
            break
        if in_section and stripped:
            section.append(line.rstrip())
    return section


def _parse_mindmap_nodes(markdown_text: str) -> List[Dict[str, Any]]:
    section_lines = _extract_mindmap_section(markdown_text)
    if not section_lines:
        return []

    nodes: List[Dict[str, Any]] = []
    stack: List[Tuple[int, str]] = []
    node_idx = 0

    for raw in section_lines:
        match = re.match(r"^(\s*)[-*]\s+(.*)$", raw)
        if not match:
            continue
        indent_spaces = len(match.group(1).replace("\t", "  "))
        depth = max(0, min(6, indent_spaces // 2))
        label = match.group(2).strip()
        if not label:
            continue

        node_idx += 1
        node_id = f"node_{node_idx}"

        while stack and stack[-1][0] >= depth:
            stack.pop()
        parent_id = stack[-1][1] if stack else None

        nodes.append(
            {
                "id": node_id,
                "label": label,
                "depth": depth,
                "parent_id": parent_id,
            }
        )
        stack.append((depth, node_id))

    return nodes


def _extract_markdown_sections(markdown_text: str) -> List[Tuple[str, List[str]]]:
    sections: List[Tuple[str, List[str]]] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []

    for raw in (markdown_text or "").splitlines():
        heading_match = re.match(r"^\s*##\s+(.+)\s*$", raw)
        if heading_match:
            if current_title:
                sections.append((current_title, current_lines))
            current_title = heading_match.group(1).strip()
            current_lines = []
            continue
        if current_title:
            current_lines.append(raw.rstrip())

    if current_title:
        sections.append((current_title, current_lines))
    return sections


def _build_fallback_mindmap_nodes(
    topic: str,
    selected_papers: List[Paper],
    markdown_text: str,
) -> List[Dict[str, Any]]:
    nodes: List[Dict[str, Any]] = []
    node_idx = 0

    def _add_node(label: str, depth: int, parent_id: Optional[str]) -> str:
        nonlocal node_idx
        text = re.sub(r"\s+", " ", (label or "").strip())
        text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
        if not text:
            text = "Untitled Node"
        node_idx += 1
        node_id = f"fallback_{node_idx}"
        nodes.append(
            {
                "id": node_id,
                "label": text[:120],
                "depth": depth,
                "parent_id": parent_id,
            }
        )
        return node_id

    root_id = _add_node(topic or "Research Mindmap", depth=0, parent_id=None)
    section_blocks = _extract_markdown_sections(markdown_text)
    excluded = {"paper links", "mindmap"}

    for section_title, section_lines in section_blocks:
        normalized_title = section_title.strip().lower()
        if normalized_title in excluded:
            continue
        parent = _add_node(section_title, depth=1, parent_id=root_id)

        bullets: List[str] = []
        for line in section_lines:
            m = re.match(r"^\s*[-*]\s+(.*)$", line)
            if not m:
                continue
            item = m.group(1).strip()
            if item:
                bullets.append(item)

        if not bullets and normalized_title == "key papers":
            bullets = [
                paper.title
                for paper in selected_papers[:4]
                if (paper.title or "").strip()
            ]

        for bullet in bullets[:4]:
            _add_node(bullet, depth=2, parent_id=parent)

        if len(nodes) >= 28:
            break

    if len(nodes) <= 1:
        keywords = _extract_keywords(
            " ".join(
                [topic]
                + [paper.title or "" for paper in selected_papers]
                + [paper.abstract or "" for paper in selected_papers]
            ),
            top_n=8,
        )
        fallback_groups = [
            ("Core Concepts", keywords[:3]),
            ("Methods", ["Baselines", "Evaluation Setup", "Ablation Study"]),
            (
                "Evidence",
                [
                    paper.title
                    for paper in selected_papers[:3]
                    if (paper.title or "").strip()
                ],
            ),
            ("Gaps", ["Missing metrics", "Generalization risk", "Reproducibility"]),
            (
                "Next Steps",
                ["Replication", "Robustness tests", "Deployment validation"],
            ),
        ]
        for group_label, items in fallback_groups:
            group_id = _add_node(group_label, depth=1, parent_id=root_id)
            for item in items[:3]:
                _add_node(str(item), depth=2, parent_id=group_id)

    return nodes


def _normalize_paper_references(text: str) -> str:
    normalized = text or ""
    # Convert [P3], P3, p3 into human-readable Paper 3 references.
    normalized = re.sub(r"\[(?:P|p)\s*(\d+)\]", r"Paper \1", normalized)
    normalized = re.sub(r"\b(?:P|p)\s*(\d+)\b", r"Paper \1", normalized)
    # Normalize accidental duplicates like "Paper Paper 3".
    normalized = re.sub(r"\bPaper\s+Paper\s+(\d+)\b", r"Paper \1", normalized)
    return normalized


def _fallback_key_insights(selected_papers: List[Paper]) -> str:
    lines: List[str] = []
    for idx, paper in enumerate(selected_papers[:5], start=1):
        abstract = (paper.abstract or "").strip()
        first_sentence = abstract.split(".")[0].strip() if abstract else ""
        if not first_sentence:
            first_sentence = (
                "Evidence is limited for detailed extraction in this workspace."
            )
        lines.append(
            f"- Insight: {first_sentence}. Why it matters: prioritizes evidence-backed planning. Evidence: Paper {idx}."
        )
    if not lines:
        lines.append(
            "- Insight: Insufficient evidence in workspace. Why it matters: add more papers before synthesis."
        )
    return "\n".join(lines)


def _ensure_report_sections(text: str, selected_papers: List[Paper]) -> str:
    content = _normalize_paper_references(text or "").strip()
    if "## Key Insights" not in content:
        content += "\n\n## Key Insights\n" + _fallback_key_insights(selected_papers)
    if "## Paper Links" not in content:
        content += "\n\n## Paper Links\n" + _paper_links_markdown(selected_papers)
    return content


def _fallback_report_markdown(
    topic: str, papers: List[Paper], depth: str = "balanced", focus_mode: str = "broad"
) -> str:
    selected = _select_ranked_papers(topic=topic, papers=papers, limit=12)
    combined_text = " ".join(
        [paper.title or "" for paper in papers]
        + [paper.abstract or "" for paper in papers]
    )
    keyword_limit = 8 if depth == "quick" else 12 if depth == "deep" else 10
    title_limit = 6 if depth == "quick" else 10 if depth == "deep" else 8
    keywords = _extract_keywords(combined_text, top_n=keyword_limit)
    representative_titles = [
        paper.title for paper in selected[:title_limit] if paper.title
    ]
    evidence_rows = []
    for idx, paper in enumerate(selected[:8], start=1):
        key_claim = ((paper.abstract or "").split(".")[0] or "").strip()
        if not key_claim:
            key_claim = "No abstract sentence available."
        evidence_rows.append(f"| Paper {idx} | {paper.title} | {key_claim[:120]} |")
    evidence_table = (
        "\n".join(
            [
                "| Paper | Title | Key claim snippet |",
                "| --- | --- | --- |",
                *evidence_rows,
            ]
        )
        if evidence_rows
        else "| Paper | Title | Key claim snippet |\n| --- | --- | --- |\n| - | - | - |"
    )

    gaps = [
        "Insufficient cross-benchmark comparability across studies.",
        "Limited reproducibility details for data preprocessing and hyperparameters.",
        "Sparse reporting on negative results and failure modes.",
        "Few studies evaluate long-term deployment constraints and robustness.",
    ]
    focus_note = {
        "methods": "Emphasis: compare methods, assumptions, and evaluation rigor.",
        "applications": "Emphasis: practical deployment value, domains, and use-case transferability.",
        "risks": "Emphasis: failure modes, uncertainty, robustness, and operational risk.",
        "broad": "Emphasis: balanced synthesis across methods, evidence, and applications.",
    }.get(
        focus_mode,
        "Emphasis: balanced synthesis across methods, evidence, and applications.",
    )

    keyword_bullets = (
        "\n".join([f"- {word.title()}" for word in keywords[:8]])
        or "- Topic not inferable from provided papers."
    )
    title_bullets = (
        "\n".join([f"- {title}" for title in representative_titles])
        or "- No representative titles available."
    )
    gap_bullets = "\n".join([f"- {gap}" for gap in gaps])
    paper_links = _paper_links_markdown(selected_papers=selected, max_items=12)
    depth_line = {
        "quick": "Report depth: quick scan for fast orientation.",
        "balanced": "Report depth: balanced review for planning and synthesis.",
        "deep": "Report depth: deep analysis for rigorous execution planning.",
    }.get(depth, "Report depth: balanced review for planning and synthesis.")

    return f"""# Research Brief: {topic}

## Executive Summary
This report summarizes {len(papers)} papers currently stored in your workspace. It captures common themes, core methods, comparative insights, and open research opportunities.
- {depth_line}
- {focus_note}

## Core Concepts
{keyword_bullets}

## Key Insights
- Insight: Model quality improves when benchmark setup is aligned across datasets. Evidence: Paper 1, Paper 2.
- Insight: Reproducibility remains a bottleneck due to incomplete implementation details. Evidence: Paper 2, Paper 3.
- Insight: Practical impact depends on domain transfer performance, not single-dataset gains. Evidence: Paper 3, Paper 4.

## Methodological Landscape
- Dominant methods rely on data-driven modeling and empirical evaluation.
- Most papers compare against prior baselines but vary in benchmark rigor.
- Evaluation metrics are often domain-specific, which reduces direct comparability.

## Comparative Findings
- Recurrent trends indicate performance gains from improved feature engineering and stronger model priors.
- Papers diverge on dataset construction assumptions and validation strategies.
- Transferability across domains remains uneven.

## Gaps and Risks
{gap_bullets}

## Future Directions
- Standardize evaluation protocols and report uncertainty with confidence intervals.
- Increase ablation depth and publish reproducible training and inference recipes.
- Validate claims on diverse and out-of-distribution settings.

## Key Papers
{title_bullets}

## Evidence Matrix
{evidence_table}

## Action Plan
- Reproduce the top two cited approaches with a single standardized benchmark suite.
- Run ablations for data preprocessing, model scaling, and objective variants.
- Track failure modes explicitly and publish robustness diagnostics.

## Paper Links
{paper_links}

## Mindmap
- {topic}
  - Core Concepts
    - {keywords[0].title() if len(keywords) > 0 else "Foundational Concepts"}
    - {keywords[1].title() if len(keywords) > 1 else "Primary Methods"}
    - {keywords[2].title() if len(keywords) > 2 else "Evaluation Focus"}
  - Methods
    - Baseline comparisons
    - Optimization and tuning
    - Validation strategies
  - Evidence
    - Representative papers
    - Quantitative outcomes
    - Reported limitations
  - Gaps
    - Reproducibility
    - Generalization risk
    - Long-term deployment
  - Next Steps
    - Better benchmarks
    - Robustness stress tests
    - Cross-domain replication
"""


def _generate_report_markdown(
    topic: str, papers: List[Paper], depth: str = "balanced", focus_mode: str = "broad"
) -> str:
    selected = _select_ranked_papers(topic=topic, papers=papers, limit=40)
    context = _build_context(topic=topic, selected_papers=selected)
    if not groq_client:
        return _fallback_report_markdown(
            topic=topic, papers=papers, depth=depth, focus_mode=focus_mode
        )

    depth_instructions = {
        "quick": "Keep sections compact (4-6 bullets each) and prioritize strongest evidence.",
        "balanced": "Provide medium depth (6-10 bullets each) with practical decisions and tradeoffs.",
        "deep": "Provide deep analysis (8-14 bullets each), include nuanced caveats and stronger synthesis.",
    }
    focus_instructions = {
        "methods": "Prioritize methodological comparisons, assumptions, benchmarks, and reproducibility.",
        "applications": "Prioritize practical applications, deployment constraints, and transferability.",
        "risks": "Prioritize risks, failure modes, uncertainty, and evidence quality concerns.",
        "broad": "Use a balanced perspective across methods, findings, applications, and risks.",
    }
    token_budget = {"quick": 2600, "balanced": 3800, "deep": 4600}.get(depth, 3800)

    prompt = (
        "Generate a rigorous, publication-grade literature analysis with a practical mindmap.\n"
        "Strict output format in markdown only.\n\n"
        "Required sections in exact order:\n"
        "# Research Brief: <topic>\n"
        "## Executive Summary\n"
        "## Core Concepts\n"
        "## Key Insights\n"
        "## Methodological Landscape\n"
        "## Comparative Findings\n"
        "## Gaps and Risks\n"
        "## Future Directions\n"
        "## Key Papers\n"
        "## Evidence Matrix\n"
        "## Action Plan\n"
        "## Paper Links\n"
        "## Mindmap\n\n"
        "Rules:\n"
        "- Keep it concise but technically deep.\n"
        "- Use bullet points for all sections except Executive Summary.\n"
        "- Every non-trivial claim must cite one or more references like 'Paper 1', 'Paper 2' from context.\n"
        "- Never use shorthand references like P1, P2, [P1], [P2].\n"
        "- In Key Insights, each bullet should include: Insight, Why it matters, Evidence (Paper references).\n"
        "- In Evidence Matrix, build a markdown table with columns: Claim | Supporting Papers | Confidence.\n"
        "- Confidence must be one of High/Medium/Low.\n"
        "- In Paper Links section, include markdown links for all cited papers in the form: Paper N: [Title](URL).\n"
        "- If URL is not available, use DOI link; if neither exists, state 'Link unavailable'.\n"
        "- Mindmap must be a hierarchical bullet tree with indentation depth up to 3 levels.\n"
        "- Base all claims on provided paper context and avoid hallucinations.\n"
        "- If evidence is weak, explicitly state 'Insufficient evidence in workspace.'\n\n"
        f"Depth profile: {depth}. {depth_instructions.get(depth, depth_instructions['balanced'])}\n"
        f"Focus profile: {focus_mode}. {focus_instructions.get(focus_mode, focus_instructions['broad'])}\n\n"
        f"Context:\n{context}"
    )

    try:
        response = groq_client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a senior research scientist and technical reviewer. "
                        "Your output must be concrete, evidence-grounded, and decision-oriented."
                    ),
                },
                {"role": "user", "content": prompt[:26000]},
            ],
            **model_config(
                task="mindmap",
                longform=True,
                max_tokens=token_budget,
                temperature=0.12 if depth == "deep" else 0.15,
            ),
        )
        text = (response.choices[0].message.content or "").strip()
        if not text:
            return _fallback_report_markdown(
                topic=topic, papers=papers, depth=depth, focus_mode=focus_mode
            )
        if not text.startswith("# Research Brief:"):
            text = f"# Research Brief: {topic}\n\n{text}"
        text = _ensure_report_sections(text=text, selected_papers=selected)
        return text[:32000]
    except Exception:
        return _fallback_report_markdown(
            topic=topic, papers=papers, depth=depth, focus_mode=focus_mode
        )


def _build_docx_bytes(markdown_text: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="DOCX export dependency missing. Install python-docx.",
        ) from exc

    document = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue

        stripped = line.lstrip(" ")
        if stripped.startswith("- "):
            indent_spaces = len(line) - len(stripped)
            paragraph = document.add_paragraph(
                stripped[2:].strip(), style="List Bullet"
            )
            if indent_spaces > 0:
                paragraph.paragraph_format.left_indent = Pt(min(indent_spaces * 4, 72))
            continue

        document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf_bytes(markdown_text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="PDF export dependency missing. Install reportlab.",
        ) from exc

    page_width, page_height = A4
    margin_x = 40
    top_y = page_height - 42
    bottom_y = 42
    line_height = 14
    max_chars = 110

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    y = top_y

    def _draw_line(text: str, font: str = "Helvetica", size: int = 11) -> None:
        nonlocal y
        if y <= bottom_y:
            pdf.showPage()
            y = top_y
        pdf.setFont(font, size)
        pdf.drawString(margin_x, y, text)
        y -= line_height

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            _draw_line("", font="Helvetica", size=11)
            continue

        font = "Helvetica"
        size = 11
        text = line

        if line.startswith("# "):
            font = "Helvetica-Bold"
            size = 15
            text = line[2:].strip()
        elif line.startswith("## "):
            font = "Helvetica-Bold"
            size = 13
            text = line[3:].strip()
        elif line.startswith("### "):
            font = "Helvetica-Bold"
            size = 12
            text = line[4:].strip()

        wrapped_lines = textwrap.wrap(text, width=max_chars) or [""]
        for wrapped in wrapped_lines:
            _draw_line(wrapped, font=font, size=size)

        if line.startswith("#"):
            y -= 3

    pdf.save()
    return buffer.getvalue()


@router.get("/", response_model=List[WorkspaceOut])
def list_workspaces(
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    return repo.list_workspaces_for_user(current_user.id)


@router.post("/", response_model=WorkspaceOut)
def create_workspace(
    payload: WorkspaceCreate,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    normalized_name = (payload.name or "").strip()
    if not normalized_name:
        raise HTTPException(status_code=400, detail="Workspace name is required")
    if len(normalized_name) > 200:
        raise HTTPException(
            status_code=422, detail="Workspace name must be 200 characters or fewer."
        )

    existing = repo.find_workspace_by_name_for_user(current_user.id, normalized_name)
    if existing:
        return existing

    return repo.create_workspace(current_user.id, normalized_name, payload.description)


def _owned_workspace_or_404(
    repo: ResearchRepository, workspace_id: int, user_id: int
) -> Workspace:
    workspace = repo.find_workspace_for_user(workspace_id, user_id)
    if not workspace:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return workspace


def _get_or_create_docspace_document(
    repo: ResearchRepository,
    current_user: User,
    workspace: Workspace,
) -> WorkspaceDocument:
    document = repo.get_docspace_document(workspace.id, current_user.id)
    if document:
        return document

    return repo.create_docspace_document(
        workspace_id=workspace.id,
        user_id=current_user.id,
        title=f"{workspace.name} Notes",
        content="",
        version=1,
    )


@router.get("/session-state", response_model=SessionStateOut)
def get_session_state(
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    state = repo.get_session_state_for_user(current_user.id)
    return _state_to_out(state)


@router.put("/session-state", response_model=SessionStateOut)
def upsert_session_state(
    payload: SessionStateUpdate,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    state = repo.get_session_state_for_user(current_user.id)
    if not state:
        state = repo.create_session_state(current_user.id)

    if payload.page_path is not None:
        state.page_path = _normalize_page_path(payload.page_path)
    if payload.workspace_id is not None:
        state.workspace_id = (
            payload.workspace_id
            if repo.workspace_exists_for_user(payload.workspace_id, current_user.id)
            else None
        )
    if payload.last_query is not None:
        state.last_query = (payload.last_query or "").strip()[:300] or None
    if payload.draft_text is not None:
        state.draft_text = (payload.draft_text or "").strip()[:12000] or None
    if payload.extra is not None:
        state.extra_json = json.dumps(payload.extra)
    state.updated_at = datetime.now(timezone.utc)

    return _state_to_out(repo.save(state))


@router.get("/{workspace_id}/docspace", response_model=DocspaceDocumentOut)
def get_docspace_document(
    workspace_id: int,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    workspace = _owned_workspace_or_404(repo, workspace_id, current_user.id)
    document = _get_or_create_docspace_document(repo, current_user, workspace)
    return DocspaceDocumentOut(
        workspace_id=workspace.id,
        title=document.title or f"{workspace.name} Notes",
        content=document.content or "",
        version=int(document.version or 1),
        updated_at=_iso_utc(document.updated_at),
    )


@router.put("/{workspace_id}/docspace", response_model=DocspaceDocumentOut)
def upsert_docspace_document(
    workspace_id: int,
    payload: DocspaceDocumentUpdate,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    workspace = _owned_workspace_or_404(repo, workspace_id, current_user.id)
    document = _get_or_create_docspace_document(repo, current_user, workspace)

    next_title = (payload.title or "").strip()
    if not next_title:
        next_title = document.title or f"{workspace.name} Notes"
    document.title = next_title[:180]
    document.content = (payload.content or "")[:200000]
    document.version = int(document.version or 1) + 1
    document.updated_at = datetime.now(timezone.utc)

    repo.save(document)
    return DocspaceDocumentOut(
        workspace_id=workspace.id,
        title=document.title or f"{workspace.name} Notes",
        content=document.content or "",
        version=int(document.version or 1),
        updated_at=_iso_utc(document.updated_at),
    )


@router.get("/{workspace_id}", response_model=WorkspaceDetail)
def get_workspace(
    workspace_id: int,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    workspace = _owned_workspace_or_404(repo, workspace_id, current_user.id)

    papers = repo.list_papers_for_workspace(workspace.id)
    chats = repo.list_chats_for_workspace(workspace.id, ascending=True)

    return WorkspaceDetail(
        id=workspace.id,
        name=workspace.name,
        description=workspace.description,
        papers=papers,
        chats=chats,
    )


@router.post("/default", response_model=WorkspaceOut)
def get_or_create_default_workspace(
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    return repo.get_or_create_default_workspace(current_user.id)


@router.delete("/{workspace_id}")
def delete_workspace(
    workspace_id: int,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    _owned_workspace_or_404(repo, workspace_id, current_user.id)
    repo.delete_workspace_graph(workspace_id)
    return {"message": "Workspace deleted successfully"}


@router.get("/{workspace_id}/files", response_model=List[WorkspaceFileOut])
def list_workspace_files(
    workspace_id: int,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    _owned_workspace_or_404(repo, workspace_id, current_user.id)
    rows = repo.list_workspace_files_for_workspace(workspace_id, current_user.id)
    return [
        WorkspaceFileOut(
            id=int(row.id or 0),
            workspace_id=int(row.workspace_id),
            user_id=int(row.user_id),
            paper_id=row.paper_id,
            kind=row.kind,
            filename=row.filename,
            storage_bucket=row.storage_bucket,
            storage_path=row.storage_path,
            download_url=row.download_url,
            content_type=row.content_type,
            size_bytes=int(row.size_bytes or 0),
            created_at=_iso_utc(getattr(row, "created_at", None)),
        )
        for row in rows
    ]


@router.get("/{workspace_id}/files/{file_id}/download")
def download_workspace_file(
    workspace_id: int,
    file_id: int,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    _owned_workspace_or_404(repo, workspace_id, current_user.id)
    row = repo.get_workspace_file_for_user(file_id, workspace_id, current_user.id)
    if not row:
        raise HTTPException(status_code=404, detail="Workspace file not found.")
    try:
        downloaded = download_bytes(storage_path=row.storage_path)
    except Exception as exc:
        raise HTTPException(
            status_code=502, detail=f"Failed to download file from storage: {str(exc)}"
        )
    headers = {"Content-Disposition": f'attachment; filename="{row.filename}"'}
    return Response(
        content=downloaded.data, media_type=downloaded.content_type, headers=headers
    )


@router.get("/{workspace_id}/export")
def export_workspace(
    workspace_id: int,
    format: str = "bibtex",
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    """Export papers in a workspace as BibTeX or CSV."""
    workspace = _owned_workspace_or_404(repo, workspace_id, current_user.id)
    papers = repo.list_papers_for_workspace(workspace.id)

    if format not in ("bibtex", "csv"):
        raise HTTPException(
            status_code=400, detail="Unsupported export format. Use 'bibtex' or 'csv'."
        )

    # CSV export
    if format == "csv":
        import csv, io, re

        def _sanitize_cell(v: str) -> str:
            s = v or ""
            # Neutralize CSV injection vectors for Excel: prefix if starts with =, +, -, @
            if s.startswith(("=", "+", "-", "@")):
                s = "'" + s
            # Remove control chars that break CSVs
            s = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", " ", s)
            return s

        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["title", "authors", "abstract", "url", "doi", "bibcode"])
        for p in papers:
            row = [
                _sanitize_cell(p.title or ""),
                _sanitize_cell(p.authors or ""),
                _sanitize_cell(p.abstract or ""),
                _sanitize_cell(p.url or ""),
                _sanitize_cell(getattr(p, "doi", "") or ""),
                _sanitize_cell(getattr(p, "bibcode", "") or ""),
            ]
            writer.writerow(row)
        content = buf.getvalue()
        filename = f"workspace-{workspace.id}.csv"
        headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
        file_record = _persist_workspace_file_if_configured(
            repo=repo,
            workspace_id=workspace.id,
            user_id=current_user.id,
            kind="workspace_export_csv",
            filename=filename,
            content=content.encode("utf-8"),
            content_type="text/csv",
        )
        if file_record:
            headers["X-Storage-Path"] = str(file_record.storage_path)
            headers["X-Storage-File-Id"] = str(file_record.id)
            headers["X-Storage-Download-Url"] = str(file_record.download_url or "")
        return Response(content=content, media_type="text/csv", headers=headers)

    # BibTeX export
    def _escape(s: str) -> str:
        return (s or "").replace("\n", " ").replace("{", "").replace("}", "").strip()

    entries = []
    for p in papers:
        key = f"paper{p.id}"
        authors = _escape(p.authors)
        title = _escape(p.title)
        year = ""
        url = _escape(p.url)
        doi = _escape(getattr(p, "doi", "") or "")
        abstract = _escape(p.abstract)
        bib_fields = []
        if doi:
            bib_fields.append(f"  doi = {{{doi}}},")
        if url:
            bib_fields.append(f"  url = {{{url}}},")
        bib_fields_str = "\n".join(bib_fields)
        optional_fields = f"\n{bib_fields_str}" if bib_fields_str else ""
        entry = (
            f"@misc{{{key},\n"
            f"  title = {{{title}}},\n"
            f"  author = {{{authors}}},{optional_fields}\n"
            f"  year = {{{year}}},\n"
            f"  abstract = {{{abstract}}}\n"
            f"}}\n"
        )
        entries.append(entry)
    content = "\n".join(entries)
    filename = f"workspace-{workspace.id}.bib"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    file_record = _persist_workspace_file_if_configured(
        repo=repo,
        workspace_id=workspace.id,
        user_id=current_user.id,
        kind="workspace_export_bibtex",
        filename=filename,
        content=content.encode("utf-8"),
        content_type="application/x-bibtex",
    )
    if file_record:
        headers["X-Storage-Path"] = str(file_record.storage_path)
        headers["X-Storage-File-Id"] = str(file_record.id)
        headers["X-Storage-Download-Url"] = str(file_record.download_url or "")
    return Response(content=content, media_type="application/x-bibtex", headers=headers)


def _resolve_research_report_inputs(
    workspace_id: int,
    payload: Optional[ResearchReportRequest],
    repo: ResearchRepository,
    current_user: User,
):
    workspace = _owned_workspace_or_404(repo, workspace_id, current_user.id)

    requested_ids = set(payload.paper_ids or []) if payload else set()
    papers = repo.list_papers_for_workspace(
        workspace.id, list(requested_ids) if requested_ids else None
    )

    if not papers:
        raise HTTPException(
            status_code=400, detail="No papers available for report generation."
        )

    raw_topic = (
        (payload.topic or workspace.name or "Research topic")
        if payload
        else (workspace.name or "Research topic")
    )
    topic = raw_topic.strip()[:160]
    depth = _normalize_report_depth(payload.depth if payload else "balanced")
    focus_mode = _normalize_focus_mode(payload.focus_mode if payload else "broad")
    selected_ids = [paper.id for paper in papers]
    return workspace, papers, topic, depth, focus_mode, selected_ids


@router.post("/{workspace_id}/research-report-preview")
def preview_research_report(
    workspace_id: int,
    payload: Optional[ResearchReportRequest] = None,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    """Generate a markdown preview of the research brief + mindmap before export."""
    workspace, papers, topic, depth, focus_mode, selected_ids = (
        _resolve_research_report_inputs(
            workspace_id=workspace_id,
            payload=payload,
            repo=repo,
            current_user=current_user,
        )
    )
    selected = _select_ranked_papers(topic=topic, papers=papers, limit=40)
    report_markdown = _generate_report_markdown(
        topic=topic,
        papers=papers,
        depth=depth,
        focus_mode=focus_mode,
    )
    mindmap_nodes = _parse_mindmap_nodes(report_markdown)
    if not mindmap_nodes:
        mindmap_nodes = _build_fallback_mindmap_nodes(
            topic=topic,
            selected_papers=selected,
            markdown_text=report_markdown,
        )
    paper_links = _paper_links_data(selected, max_items=24)
    return {
        "workspace_id": workspace.id,
        "workspace_name": workspace.name,
        "topic": topic,
        "depth": depth,
        "focus_mode": focus_mode,
        "paper_count": len(papers),
        "selected_paper_ids": selected_ids,
        "generated_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "markdown": report_markdown,
        "mindmap_nodes": mindmap_nodes,
        "paper_links": paper_links,
    }


@router.post("/{workspace_id}/research-report")
def export_research_report(
    workspace_id: int,
    format: str = "pdf",
    payload: Optional[ResearchReportRequest] = None,
    repo: ResearchRepository = Depends(get_research_repository),
    current_user: User = Depends(get_current_user),
):
    """
    Generate a structured research brief + mindmap from workspace papers.
    Export options: PDF or DOCX.
    """
    if format not in ("pdf", "docx"):
        raise HTTPException(
            status_code=400, detail="Unsupported export format. Use 'pdf' or 'docx'."
        )

    workspace, papers, topic, depth, focus_mode, _selected_ids = (
        _resolve_research_report_inputs(
            workspace_id=workspace_id,
            payload=payload,
            repo=repo,
            current_user=current_user,
        )
    )
    report_markdown = _generate_report_markdown(
        topic=topic,
        papers=papers,
        depth=depth,
        focus_mode=focus_mode,
    )

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    safe_topic = _safe_filename(topic)

    if format == "docx":
        content = _build_docx_bytes(report_markdown)
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = f"{safe_topic}-mindmap-{timestamp}.docx"
    else:
        content = _build_pdf_bytes(report_markdown)
        media_type = "application/pdf"
        filename = f"{safe_topic}-mindmap-{timestamp}.pdf"

    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    file_record = _persist_workspace_file_if_configured(
        repo=repo,
        workspace_id=workspace.id,
        user_id=current_user.id,
        kind=f"research_report_{format}",
        filename=filename,
        content=content,
        content_type=media_type,
    )
    if file_record:
        headers["X-Storage-Path"] = str(file_record.storage_path)
        headers["X-Storage-File-Id"] = str(file_record.id)
        headers["X-Storage-Download-Url"] = str(file_record.download_url or "")
    return Response(content=content, media_type=media_type, headers=headers)
